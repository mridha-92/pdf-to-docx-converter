"""High-precision PDF-to-DOCX conversion engine.

Pipeline:
  1. analyze_pdf()  – detect encryption, corruption, and whether the PDF
                      is text-based (native) or scanned (image-only).
  2. convert_pdf_to_docx() – prefer pdf2docx for native documents (keeps
                      tables and multi-column layout); on failure, fall back
                      to a span-level formatter built on PyMuPDF
                      get_text("dict") + python-docx that maps fonts,
                      bold/italic flags, sizes and colors onto Run objects.
  3. Graceful exceptions are raised for corrupted / encrypted / scanned
                      documents instead of crashing.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF

logger = logging.getLogger("converter")


# ---------------------------------------------------------------------------
# Domain errors
# ---------------------------------------------------------------------------
class PDFConversionError(Exception):
    """Base error for all conversion failures."""


class PDFCorruptedError(PDFConversionError):
    """The PDF is truncated or unreadable."""


class PDFEncryptedError(PDFConversionError):
    """The PDF is password-protected and cannot be read."""


class PDFScannedError(PDFConversionError):
    """The PDF contains no extractable text layer (scanned images only)."""


# ---------------------------------------------------------------------------
# File analysis
# ---------------------------------------------------------------------------
@dataclass
class PDFAnalysis:
    path: Path
    page_count: int
    text_based: bool          # True = at least one page has extractable text
    encrypted: bool
    has_text_pages: int       # pages that contain any extractable text


def analyze_pdf(pdf_path: Path) -> PDFAnalysis:
    """Inspect a PDF and classify it (encrypted / corrupted / text-based)."""
    if not pdf_path.exists():
        raise PDFConversionError("The uploaded file could not be found.")

    try:
        doc = fitz.open(pdf_path)
    except Exception as exc:  # noqa: BLE001 - PyMuPDF raises many error types
        raise PDFCorruptedError(
            "The file is not a valid PDF or is corrupted."
        ) from exc

    try:
        if doc.needs_pass:
            return PDFAnalysis(
                path=Path(pdf_path),
                page_count=0,
                text_based=False,
                encrypted=True,
                has_text_pages=0,
            )

        page_count = doc.page_count
        has_text_pages = 0
        for page in doc:
            text = page.get_text("text").strip()
            if text:
                has_text_pages += 1
        text_based = has_text_pages > 0
    finally:
        if not doc.is_closed:
            doc.close()

    return PDFAnalysis(
        path=Path(pdf_path),
        page_count=page_count,
        text_based=text_based,
        encrypted=False,
        has_text_pages=has_text_pages,
    )


# ---------------------------------------------------------------------------
# Native conversion via pdf2docx
# ---------------------------------------------------------------------------
def _count_source_lines(pdf_path: Path) -> int:
    """Count extractable text lines in the PDF (used to verify output)."""
    count = 0
    with fitz.open(pdf_path) as pdf:
        for page in pdf:
            count += len(page.get_text("dict").get("blocks", []))
    return count


def _convert_native(pdf_path: Path, docx_path: Path) -> None:
    """Convert using pdf2docx, which preserves tables and multi-column flow.

    pdf2docx swallows per-page errors and can silently emit an empty/partial
    document, so the result is verified afterwards and raised to trigger the
    span-level fallback when content is missing.
    """
    from docx import Document
    from pdf2docx import Converter

    cv = Converter(str(pdf_path))
    try:
        cv.convert(str(docx_path))
    finally:
        cv.close()

    # Verify the result contains meaningful content.
    try:
        result = Document(docx_path)
        paragraphs = [p.text for p in result.paragraphs if p.text.strip()]
        table_count = len(result.tables)
        total = len(paragraphs) + table_count
    except Exception as exc:  # noqa: BLE001
        raise PDFConversionError(
            "The PDF could not be converted (empty result)."
        ) from exc

    source_lines = max(_count_source_lines(pdf_path), 1)
    if total == 0 or total < source_lines * 0.4:
        raise PDFConversionError(
            "pdf2docx produced an incomplete result; switching engines."
        )


# ---------------------------------------------------------------------------
# Precision span-level fallback (PyMuPDF -> python-docx)
# ---------------------------------------------------------------------------
# PyMuPDF text flags
_FLAG_SUPERSCRIPTED = 2**0
_FLAG_ITALIC = 2**1
_FLAG_MONOSPACED = 2**3
_FLAG_BOLD = 2**4

# Map common PDF base-14 fonts to Windows-safe Word fonts.
_FONT_MAP = {
    "helvetica": "Arial",
    "helvetica-bold": "Arial",
    "helvetica-oblique": "Arial",
    "times-roman": "Times New Roman",
    "times-bold": "Times New Roman",
    "times-italic": "Times New Roman",
    "times-bolditalic": "Times New Roman",
    "courier": "Courier New",
    "courier-bold": "Courier New",
    "courier-oblique": "Courier New",
    "courier-boldoblique": "Courier New",
    "symbol": "Symbol",
    "zapfdingbats": "Wingdings",
}


def _resolve_font(font_name: str) -> str:
    """Translate a PDF font name into a python-docx font name."""
    if not font_name:
        return "Arial"
    base = font_name.lower().split("+")[-1].replace("-", "").replace("_", "")
    if base in _FONT_MAP:
        return _FONT_MAP[base]
    # Fallback: strip subsetting prefix, return best-effort name.
    return font_name.split("+")[-1] or "Arial"


def _font_size_to_heading(size: float, median: float) -> str:
    """Heuristic: relative font size -> Word heading style."""
    if size >= max(20.0, median * 1.9):
        return "Heading 1"
    if size >= max(16.0, median * 1.5):
        return "Heading 2"
    if size >= max(13.0, median * 1.2):
        return "Heading 3"
    return "Normal"


def _convert_fallback(pdf_path: Path, docx_path: Path) -> None:
    """Span-level fallback converter for complex/native documents."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Consistent base style so headings/body look clean.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    all_sizes: list[float] = []

    with fitz.open(pdf_path) as pdf:
        for page in pdf:
            data = page.get_text("dict")
            for block in data.get("blocks", []):
                if block.get("type") != 0:  # skip images for now
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if text.strip():
                            all_sizes.append(span.get("size", 0.0))

        median = _median(all_sizes) if all_sizes else 11.0

        for page in pdf:
            data = page.get_text("dict")
            for block in data.get("blocks", []):
                if block.get("type") != 0:
                    continue
                # Merge the line's spans into one paragraph per line so that
                # bold/italic runs within a single line are preserved.
                for line in block.get("lines", []):
                    spans = [s for s in line.get("spans", []) if s.get("text", "").strip()]
                    if not spans:
                        continue

                    first = spans[0]
                    heading = _font_size_to_heading(first.get("size", 11.0), median)

                    paragraph = doc.add_paragraph(style=heading)
                    for span in spans:
                        text = span.get("text", "")
                        if not text:
                            continue

                        # Split multi-line span content into separate runs.
                        parts = text.split("\n")
                        for idx, part in enumerate(parts):
                            if part:
                                run = paragraph.add_run(part)
                                run.bold = bool(span.get("flags", 0) & _FLAG_BOLD)
                                run.italic = bool(span.get("flags", 0) & _FLAG_ITALIC)
                                run.font.name = _resolve_font(span.get("font", ""))
                                run.font.size = Pt(max(span.get("size", 11.0) * 0.75, 8.0))
                                superscripted = bool(span.get("flags", 0) & _FLAG_SUPERSCRIPTED)
                                if superscripted:
                                    run.font.superscript = True
                                color = span.get("color")
                                if color and color != 0 and (color & 0xFFFFFF) != 0x000000:
                                    run.font.color.rgb = RGBColor(
                                        (color >> 16) & 0xFF,
                                        (color >> 8) & 0xFF,
                                        color & 0xFF,
                                    )
                            # Blank segment between \n => line break.
                            if idx < len(parts) - 1:
                                paragraph.add_run().add_break()

    doc.save(docx_path)


def _median(values: list[float]) -> float:
    ordered = sorted(values)
    mid = len(ordered) // 2
    if len(ordered) % 2 == 0:
        return (ordered[mid - 1] + ordered[mid]) / 2
    return ordered[mid]


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------
def convert_pdf_to_docx(pdf_path: Path, docx_path: Path) -> dict[str, Any]:
    """Convert a PDF to DOCX, returning metadata about the result.

    Raises PDFConversionError subclasses with user-friendly messages.
    """
    analysis = analyze_pdf(pdf_path)

    if analysis.encrypted:
        raise PDFEncryptedError(
            "This PDF is password-protected. Remove the password and try again."
        )
    if analysis.page_count == 0:
        raise PDFConversionError("The PDF contains no pages.")
    if not analysis.text_based:
        raise PDFScannedError(
            "This PDF appears to be scanned (no extractable text). OCR is not "
            "available in this build."
        )

    try:
        _convert_native(pdf_path, docx_path)
        engine = "pdf2docx"
    except Exception as exc:  # noqa: BLE001 - fall back on any failure
        logger.warning("pdf2docx failed (%s); using span-level fallback.", exc)
        try:
            _convert_fallback(pdf_path, docx_path)
            engine = "span-fallback"
        except PDFConversionError:
            raise
        except Exception as exc2:  # noqa: BLE001
            raise PDFConversionError(
                "Conversion failed while processing the document."
            ) from exc2

    return {
        "page_count": analysis.page_count,
        "text_pages": analysis.has_text_pages,
        "engine": engine,
    }
